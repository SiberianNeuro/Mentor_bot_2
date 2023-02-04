import os

from typing import Union, Optional, NamedTuple
from datetime import datetime

import docx
from docx.opc.exceptions import PackageNotFoundError

from aiogram import Bot, types
import sqlalchemy as sa
from sqlalchemy.orm import sessionmaker

from app.models.user import User


class Protocol(NamedTuple):
    fullname: str
    user: Union[int, str]
    stage_id: int
    result_id: int
    score: float
    exam_date: datetime.date
    retake_date: Union[datetime.date, None]


# Парсер для протоколов опроса
async def file_parser(document: types.document.Document, bot: Bot, db: sessionmaker) -> Protocol | str:
    """

    :param db: sqlalchemy session object
    :param bot: telegram bot object
    :param document: telegram file-id
    :return: тапл из полного имени пользователя, айди формата опроса, айди результата опроса,
    баллы за опрос, дата опроса и дата пересдачи
    """

    file_id = await bot.get_file(document.file_id)
    await bot.download_file(file_id.file_path, document.file_name)
    try:
        d = docx.Document(document.file_name)
    except PackageNotFoundError:
        os.remove(document.file_name)
        return "Не распознал расширение протокола. Я кушаю протоколы только в формате .docx"

    general_table = d.tables[0]
    retake_date = None
    if "Призывные мероприятия" in d.paragraphs[3].text:
        stage_id = 1
    elif "в середине теоретического цикла" in d.paragraphs[3].text:
        stage_id = 2
    elif "при аттестации" in d.paragraphs[3].text:
        stage_id = 3
    elif "при переводе" in d.paragraphs[3].text:
        stage_id = 4
    elif "при экзаменации" in d.paragraphs[3].text:
        stage_id = 5
    else:
        os.remove(document.file_name)
        return "Я не распознал протокол.\n\nПроверь, актуальным ли протоколом ты пользуешься.\n" \
               "Если нет, то узнай у руководителя актуальную версю, перепиши результаты туда, и снова " \
               "отправь мне."

    scores_table = d.tables[2] if stage_id == 3 else d.tables[1]
    results_table = d.tables[3] if stage_id == 3 else d.tables[2]
    general = []
    scores = []
    results = []

    for r in general_table.rows:
        general.append([cell.text for cell in r.cells])
    for r in scores_table.rows:
        scores.append([cell.text for cell in r.cells])
    for r in results_table.rows:
        results.append([cell.text for cell in r.cells])

    fullname = general[0][1].strip()
    async with db.begin() as ses:
        sql = sa.select(User.id).where(sa.and_(User.fullname == fullname, User.active == 1))
        result = await ses.execute(sql)
        user = result.scalars().first()
    if not user:
        os.remove(document.file_name)
        return 'Я не нашел такого стажера в базе данных. Проверь, пожалуйста, правильно ли заполнено ФИО.\n\n' \
               'Если все правильно, уточни у стажера, проходил ли он у меня регистрацию.'
    if stage_id == 4:
        score = 0.0
    else:
        score = scores[-1][2]
        if score == '':
            os.remove(document.file_name)
            return "Я не нашел итоговое количество баллов.\n\nВнимательно посмотри, заполнил ли ты их. " \
                   "Если не заполнил - заполняй и присылай мне протокол повторно."

    if stage_id == 1:
        result_id = 3
    else:
        if results[-1][2] != "":
            result_id = 1
        elif results[-2][2] != "":
            result_id = 2
            try:
                retake_date = datetime.strptime(results[-2][3], "%d.%m.%Y")
            except ValueError:
                os.remove(document.file_name)
                return "Неверно заполнено поле даты переаттестации.\n\nПожалуйста, введи дату в формате " \
                       "<i>ДД.ММ.ГГГГ</i> и пришли мне повторно."
        else:
            result_id = 3

    if general[-1][1] != '':
        try:
            exam_date = datetime.strptime(general[-1][1], "%d.%m.%Y")
        except ValueError:
            os.remove(document.file_name)
            return "Поле 'Дата проведения проф.опроса' заполнено некорректно.\n\nПожалуйста, введи дату в формате " \
                   "<i>ДД.ММ.ГГГГ</i> и пришли мне повторно."
    else:
        os.remove(document.file_name)
        return "Не заполнено поле 'Дата проведения проф.опроса'.\n\nЗаполни его, пожалуйста, и пришли мне снова."

    try:
        score = float(score.replace(',', '.'))
    except AttributeError:
        score = float(score)
    finally:
        os.remove(document.file_name)

        return Protocol(
            fullname=fullname,
            user=user,
            stage_id=stage_id,
            result_id=result_id,
            score=score,
            exam_date=exam_date,
            retake_date=retake_date
        )
